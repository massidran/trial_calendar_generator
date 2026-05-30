import datetime as dt
import docx
import csv
import sys
import os

COURT_HOLIDAYS = set()

def resource_path(filename: str) -> str:
    if getattr(sys, "frozen", False):
        base_path = sys._MEIPASS
    else:
        base_path = os.path.dirname(os.path.abspath(__file__))
    return os.path.join(base_path, filename)

def load_holidays():
    holidays_path = resource_path("holidays.csv")
    with open(holidays_path, "r", newline="") as f:
        reader = csv.reader(f)
        for row in reader:
            if not row:
                continue
            COURT_HOLIDAYS.add(dt.date(int(row[0]), int(row[1]), int(row[2])))

def is_court_holiday(date: dt.date) -> bool:
    return date.date() in COURT_HOLIDAYS

def sub_days(date: dt.date, days: int) -> dt.datetime:
    date -= dt.timedelta(days=days)
    while date.weekday() >= 5 or is_court_holiday(date):
        date += dt.timedelta(days=1)
    return date

def sub_court_days(date: dt.date, days: int) -> dt.datetime:
    count = 0
    while count < days:
        date -= dt.timedelta(days=1)
        if date.weekday() < 5 and not is_court_holiday(date):
            count += 1
    return date

def open_doc():
    doc_path = resource_path("trial_calendar_template.docx")
    return docx.Document(doc_path)

def find_and_replace_head(doc: docx.Document, placeholder: str, replacement: str):
    for p in doc.paragraphs:
        if placeholder in p.text:
            for run in p.runs:
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, replacement)
                    run.font.name = "Times New Roman"
                    run.font.size = docx.shared.Pt(24)
                    run.bold = True
                    return

def find_and_replace(doc: docx.Document, placeholder: str, replacement: str):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if placeholder in p.text:
                        p.text = p.text.replace(placeholder, replacement)
                        for run in p.runs:
                            run.font.name = "Cambria"
                            run.bold = True
                        return